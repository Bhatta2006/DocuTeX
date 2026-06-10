import re
import os
import sys
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Global state tracking for paragraph formatting
just_after_heading: bool = False
just_after_table_fig: bool = False

# XML helpers for tables
def set_cell_border(cell: Any, **kwargs: Any) -> None:
    """
    Set cell borders.
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000"},
        bottom={"sz": 6, "val": "single", "color": "000000"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        # Clear existing border of this type if any
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
            
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn(f'w:{key}'), str(val))
            tcBorders.append(element)
        else:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'none')
            tcBorders.append(element)

def set_cell_margins(cell: Any, top: int = 100, bottom: int = 100, left: int = 150, right: int = 150) -> None:
    """Set cell margins in dxa (1 pt = 20 dxa)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# LaTeX parsing helpers
def get_nested_brackets(text: str, start_idx: int) -> Tuple[str, int]:
    """
    Find matching curly braces starting at start_idx (which must be a '{').
    Returns (matched_substring, next_index).
    """
    if start_idx >= len(text) or text[start_idx] != '{':
        start_idx = text.find('{', start_idx)
        if start_idx == -1:
            return "", len(text)
            
    depth = 0
    content = []
    for idx in range(start_idx, len(text)):
        char = text[idx]
        if char == '{':
            depth += 1
        elif char == '}':
            depth -= 1
        content.append(char)
        if depth == 0:
            return "".join(content), idx + 1
    return "".join(content), len(text)

def get_optional_brackets(text: str, start_idx: int) -> Tuple[str, int]:
    """
    Find matching square brackets starting at start_idx (which must be a '[').
    Returns (matched_substring, next_index).
    """
    if start_idx >= len(text) or text[start_idx] != '[':
        start_idx = text.find('[', start_idx)
        if start_idx == -1:
            return "", start_idx
            
    depth = 0
    content = []
    for idx in range(start_idx, len(text)):
        char = text[idx]
        if char == '[':
            depth += 1
        elif char == ']':
            depth -= 1
        content.append(char)
        if depth == 0:
            return "".join(content), idx + 1
    return "".join(content), len(text)

def extract_command_arg(text: str, cmd_name: str, start_pos: int = 0) -> Tuple[Optional[str], int]:
    """
    Searches for \cmd_name{arg} and returns (arg, end_pos).
    """
    pattern = re.escape(cmd_name) + r'\s*\{'
    match = re.search(pattern, text[start_pos:])
    if not match:
        return None, start_pos
    brace_start = start_pos + match.end() - 1
    content, end_pos = get_nested_brackets(text, brace_start)
    return content[1:-1].strip(), end_pos

def extract_consecutive_args(text: str, start_pos: int, count: int) -> Tuple[List[str], int]:
    """
    Extract count consecutive curly-braced arguments starting from start_pos.
    Returns (list_of_args, end_pos).
    """
    args = []
    pos = start_pos
    for _ in range(count):
        match = re.match(r'\s*\{', text[pos:])
        if not match:
            break
        brace_start = pos + match.end() - 1
        content_brace, end_pos = get_nested_brackets(text, brace_start)
        args.append(content_brace[1:-1].strip())
        pos = end_pos
    return args, pos

def strip_latex_command(text: str, cmd_name: str) -> str:
    """
    Finds and strips a LaTeX command and its nested curly-braced argument from the text.
    """
    pattern = re.escape(cmd_name) + r'\s*\{'
    while True:
        match = re.search(pattern, text)
        if not match:
            break
        brace_start = match.end() - 1
        _, end_pos = get_nested_brackets(text, brace_start)
        text = text[:match.start()] + text[end_pos:]
    return text

def preprocess_latex(text: str) -> str:
    """
    Preprocess LaTeX content to replace custom resume templates and command shortcuts.
    """
    try:
        # 1. Clean up spacing commands and helpers
        text = text.replace(r'\quad', ' | ')
        text = text.replace(r'\qquad', ' | ')
        text = text.replace(r'\hfill', '   ')
        
        # 2. Simple command replacements for resume environments
        text = text.replace(r'\resumeSubHeadingListStart', '')
        text = text.replace(r'\resumeSubHeadingListEnd', '')
        text = text.replace(r'\resumeItemListStart', r'\begin{itemize}')
        text = text.replace(r'\resumeItemListEnd', r'\end{itemize}')
        
        # 3. Replace custom list items \resumeItem{content} -> \item content
        while True:
            match = re.search(r'\\resumeItem\s*\{', text)
            if not match:
                break
            brace_start = match.end() - 1
            content_brace, end_pos = get_nested_brackets(text, brace_start)
            content = content_brace[1:-1].strip()
            text = text[:match.start()] + f"\\item {content}" + text[end_pos:]
            
        # 4. Replace \resumeSubheading{company}{loc}{role}{date}
        while True:
            match = re.search(r'\\resumeSubheading\s*', text)
            if not match:
                break
            args, end_pos = extract_consecutive_args(text, match.end(), 4)
            if len(args) == 4:
                company, loc, role, date = args
                replacement = f"\\subsection{{{company} -- {role}}}\n\\textit{{{date} | {loc}}}\n\n"
                text = text[:match.start()] + replacement + text[end_pos:]
            else:
                # Fallback if args aren't 4
                text = text[:match.start()] + text[end_pos:]
                
        # 5. Replace \resumeProjectHeading{project}{date}
        while True:
            match = re.search(r'\\resumeProjectHeading\s*', text)
            if not match:
                break
            args, end_pos = extract_consecutive_args(text, match.end(), 2)
            if len(args) == 2:
                project, date = args
                replacement = f"\\subsection{{{project}}}\n\\textit{{{date}}}\n\n"
                text = text[:match.start()] + replacement + text[end_pos:]
            else:
                text = text[:match.start()] + text[end_pos:]
                
        # 6. Parse \href{url}{text} and replace with text
        while True:
            match = re.search(r'\\href\s*\{', text)
            if not match:
                break
            url_start = match.end() - 1
            url_brace, url_end = get_nested_brackets(text, url_start)
            url = url_brace[1:-1]
            
            text_match = re.match(r'\s*\{', text[url_end:])
            if not text_match:
                text = text[:match.start()] + url + text[url_end:]
                continue
            text_start = url_end + text_match.end() - 1
            text_brace, text_end = get_nested_brackets(text, text_start)
            display_text = text_brace[1:-1]
            text = text[:match.start()] + display_text + text[text_end:]
            
        return text
    except Exception as e:
        print(f"Warning: LaTeX pre-processing encountered an error: {e}")
        return text


def replace_fractions(m_str: str) -> str:
    """Replace LaTeX fractions \\frac{a}{b} or shorthands with (a)/(b) recursively."""
    # 1. Replace \frac\partial f_x\partial x with (\partial f_x)/(\partial x)
    m_str = re.sub(r'\\frac\\partial\s*f_([a-zA-Z])\\partial\s*([a-zA-Z])', r'(\\partial f_\1)/(\\partial \2)', m_str)
    
    # 2. Replace \frac1W with (1)/(W)
    m_str = re.sub(r'\\frac\s*([0-9a-zA-Z])\s*([0-9a-zA-Z])', r'(\1)/(\2)', m_str)
    
    # 3. Standard recursive braced fraction replacement
    while True:
        match = re.search(r'\\frac\s*\{', m_str)
        if not match:
            break
        num_start = match.end() - 1
        num_brace, num_end = get_nested_brackets(m_str, num_start)
        num = num_brace[1:-1]
        
        den_match = re.match(r'\s*\{', m_str[num_end:])
        if not den_match:
            break
        den_start = num_end + den_match.end() - 1
        den_brace, den_end = get_nested_brackets(m_str, den_start)
        den = den_brace[1:-1]
        
        m_str = m_str[:match.start()] + f"({num})/({den})" + m_str[den_end:]
    return m_str

def translate_math_symbols(m: str) -> str:
    """Translates LaTeX math commands and Greek letters into Unicode equivalents."""
    # Greek letters
    greek_letters = {
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ', r'\epsilon': 'ε',
        r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ', r'\iota': 'ι', r'\kappa': 'κ',
        r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π',
        r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ', r'\upsilon': 'υ', r'\phi': 'φ',
        r'\chi': 'χ', r'\psi': 'ψ', r'\omega': 'ω',
        r'\Sigma': 'Σ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ', r'\Xi': 'Ξ',
        r'\Pi': 'Π', r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω', r'\Gamma': 'Γ',
    }
    # Mathematical operators
    math_ops = {
        r'\le': '≤', r'\ge': '≥', r'\pm': '±', r'\cdot': '·', r'\times': '×',
        r'\div': '÷', r'\neq': '≠', r'\approx': '≈', r'\propto': '∝', r'\infty': '∞',
        r'\partial': '∂', r'\nabla': '∇', r'\sum': '∑', r'\prod': '∏', r'\int': '∫',
        r'\sqrt': '√', r'\in': '∈', r'\notin': '∉', r'\subset': '⊆', r'\cup': '∪',
        r'\cap': '∩', r'\dots': '...', r'\ldots': '...', r'\cdots': '···',
    }
    
    # Sort keys by length in descending order to avoid partial matching (e.g. \sum vs \subset)
    all_symbols = {**greek_letters, **math_ops}
    sorted_keys = sorted(all_symbols.keys(), key=len, reverse=True)
    
    for key in sorted_keys:
        m = m.replace(key, all_symbols[key])
        
    return m

def latex_math_to_tagged_string(math_str: str) -> str:
    """Converts a math equation string to a tagged Unicode markup string."""
    try:
        m = math_str.strip()
        if m.startswith('$') and m.endswith('$'):
            m = m[1:-1].strip()
            
        # 1. Remove formatting commands first to prevent partial translation of keywords like \left (containing \le)
        m = re.sub(r'\\(mathbf|mathrm|mathit|mathsf|mathtt|left|right)', '', m)
        m = m.replace(r'\{', '{').replace(r'\}', '}')
        
        # 2. Replace fractions
        m = replace_fractions(m)
        
        # 3. Handle hats, bars, vectors
        accents = {
            r'\hat': '̂',
            r'\bar': '̄',
            r'\tilde': '̃',
            r'\vec': '⃗',
        }
        for acc_cmd, unicode_acc in accents.items():
            while True:
                match = re.search(re.escape(acc_cmd) + r'\s*\{', m)
                if not match:
                    break
                brace_start = match.end() - 1
                content_brace, end_pos = get_nested_brackets(m, brace_start)
                content = content_brace[1:-1]
                content_translated = translate_math_symbols(content)
                m = m[:match.start()] + content_translated + unicode_acc + m[end_pos:]
                
        # 4. Translate math symbols
        m = translate_math_symbols(m)
        
        # 5. Replace \text{...} with <normal>...</normal>
        while True:
            match = re.search(re.escape(r'\text') + r'\s*\{', m)
            if not match:
                break
            brace_start = match.end() - 1
            content_brace, end_pos = get_nested_brackets(m, brace_start)
            content = content_brace[1:-1]
            m = m[:match.start()] + f"<normal>{content}</normal>" + m[end_pos:]
            
        # Translate subscripts and superscripts
        while True:
            match = re.search(r'_\s*\{', m)
            if not match:
                break
            brace_start = match.end() - 1
            content_brace, end_pos = get_nested_brackets(m, brace_start)
            content = content_brace[1:-1]
            m = m[:match.start()] + f"<sub>{content}</sub>" + m[end_pos:]
            
        m = re.sub(r'_([a-zA-Z0-9])', r'<sub>\1</sub>', m)
        
        while True:
            match = re.search(r'\^\s*\{', m)
            if not match:
                break
            brace_start = match.end() - 1
            content_brace, end_pos = get_nested_brackets(m, brace_start)
            content = content_brace[1:-1]
            m = m[:match.start()] + f"<sup>{content}</sup>" + m[end_pos:]
            
        m = re.sub(r'\^([a-zA-Z0-9])', r'<sup>\1</sup>', m)
        
        # Italicize variables
        parts = re.split(r'(</?\w+>)', m)
        active_normal = False
        
        for idx, part in enumerate(parts):
            if not part:
                continue
            if part == '<normal>':
                active_normal = True
            elif part == '</normal>':
                active_normal = False
            elif not part.startswith('<'):
                if not active_normal:
                    words = re.split(r'([^a-zA-Z])', part)
                    for w_idx, word in enumerate(words):
                        if word.isalpha() and word not in ['sin', 'cos', 'tan', 'log', 'ln', 'exp', 'lim', 'max', 'min', 'det', 'Var', 'Div', 'atan', 'atan2', 'in', 'TOTAL', 'Total']:
                            words[w_idx] = f"<italic>{word}</italic>"
                    part = "".join(words)
                    
                    # Greek characters
                    greek_chars = 'αβγδεζηθικλμνξοπρστυφχψωΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ'
                    for g in greek_chars:
                        part = part.replace(g, f'<italic>{g}</italic>')
                parts[idx] = part
                
        m = "".join(parts)
        return m
    except Exception as e:
        print(f"Warning: Failed to parse math formula: {math_str}. Error: {e}")
        return math_str

def resolve_citations(text: str, citation_map: Dict[str, int]) -> str:
    """Resolves \\cite{...} command into citation numbers."""
    def cite_replacer(match: Any) -> str:
        labels = [l.strip() for l in match.group(1).split(',')]
        resolved = []
        for label in labels:
            if label in citation_map:
                resolved.append(str(citation_map[label]))
            else:
                resolved.append(label)
        return "[" + ", ".join(resolved) + "]"
        
    text = re.sub(r'\\cite\{([^}]+)\}', cite_replacer, text)
    return text

def clean_latex_text(text: str, label_map: Dict[str, Any], citation_map: Dict[str, int]) -> str:
    """Translates LaTeX markup structures to Tagged text format and unescapes symbols."""
    try:
        # Strip comments
        text = re.sub(r'%.*?\n', '\n', text)
        
        # Convert inline math $...$
        text = re.sub(r'\$([^$]+)\$', lambda m: latex_math_to_tagged_string(m.group(1)), text)
        
        # Common text commands
        cmds = [
            (r'\textit', '<italic>', '</italic>'),
            (r'\textbf', '<bold>', '</bold>'),
            (r'\texttt', '<monospace>', '</monospace>'),
            (r'\textsuperscript', '<sup>', '</sup>'),
            (r'\textsubscript', '<sub>', '</sub>'),
        ]
        
        for cmd_name, open_tag, close_tag in cmds:
            while True:
                match = re.search(re.escape(cmd_name) + r'\s*\{', text)
                if not match:
                    break
                brace_start = match.end() - 1
                content_brace, end_pos = get_nested_brackets(text, brace_start)
                content = content_brace[1:-1]
                text = text[:match.start()] + open_tag + content + close_tag + text[end_pos:]
                
        # Resolve cross-references and citations
        text = resolve_citations(text, citation_map)
        
        def ref_replacer(match: Any) -> str:
            label = match.group(1).strip()
            if label in label_map:
                return str(label_map[label])
            return label
            
        text = re.sub(r'\\ref\{([^}]+)\}', ref_replacer, text)
        
        # Remove outstanding label declarations
        text = re.sub(r'\\label\{[^}]+\}', '', text)
        
        # Translate double-backslashes (forced line breaks) to soft line breaks
        text = text.replace(r'\\', '\n')
        
        # Unescape characters from LaTeX
        text = text.replace(r'\_', '_')
        text = text.replace(r'\&', '&')
        text = text.replace(r'\%', '%')
        text = text.replace(r'\#', '#')
        text = text.replace(r'\$', '$')
        text = text.replace('~', ' ')
        
        # Strip formatting command words
        text = re.sub(r'\\(rowcolor|cellcolor|columncolor)\s*\{[^}]*\}', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\\color\s*\{[^}]*\}', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\\arraybackslash', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\\(Huge|huge|Large|large|normalsize|small|footnotesized|scriptsize|tiny)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\\(scshape|bfseries|itshape|slshape|sffamily|ttfamily|rmfamily|selectfont)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\\(vspace|hspace)\*?\{[^}]*\}', '', text)
        text = re.sub(r'\\(quad|qquad)', ' | ', text)
        text = re.sub(r'\\(hfill|vfill|noindent)', ' ', text)

        # Remove any unescaped curly braces by temporarily saving escaped ones
        text = text.replace(r'\{', 'TEMP_OPEN_BRACE').replace(r'\}', 'TEMP_CLOSE_BRACE')
        text = text.replace('{', '').replace('}', '')
        text = text.replace('TEMP_OPEN_BRACE', '{').replace('TEMP_CLOSE_BRACE', '}')

        return text
    except Exception as e:
        print(f"Warning: Failed to clean LaTeX text: {text[:100]}... Error: {e}")
        return text


def add_formatted_text(p: Any, text: str, font_size: int = 10, is_abstract: bool = False) -> None:
    """Processes tagged text and appends formatted runs to the paragraph."""
    tokens = re.split(r'(</?\w+>)', text)
    
    active_italic = False
    active_bold = is_abstract
    active_mono = False
    active_sub = False
    active_super = False
    
    for token in tokens:
        if not token:
            continue
        if token == '<italic>':
            active_italic = True
        elif token == '</italic>':
            active_italic = False
        elif token == '<bold>':
            active_bold = True
        elif token == '</bold>':
            active_bold = is_abstract
        elif token == '<monospace>':
            active_mono = True
        elif token == '</monospace>':
            active_mono = False
        elif token == '<sub>':
            active_sub = True
        elif token == '</sub>':
            active_sub = False
        elif token == '<sup>':
            active_super = True
        elif token == '</sup>':
            active_super = False
        elif token.startswith('<') and token.endswith('>'):
            pass
        else:
            run = p.add_run(token)
            run.font.name = 'Courier New' if active_mono else 'Times New Roman'
            run.font.size = Pt(font_size)
            run.font.italic = active_italic
            run.font.bold = active_bold
            if active_sub:
                run.font.subscript = True
            if active_super:
                run.font.superscript = True

def add_math_to_paragraph(p: Any, math_text: str) -> None:
    """Appends formatted math runs to a paragraph based on math tags."""
    tokens = re.split(r'(</?\w+>)', math_text)
    
    active_italic = False
    active_bold = False
    active_sub = False
    active_super = False
    active_normal = False
    
    for token in tokens:
        if not token:
            continue
        if token == '<italic>':
            active_italic = True
        elif token == '</italic>':
            active_italic = False
        elif token == '<bold>':
            active_bold = True
        elif token == '</bold>':
            active_bold = False
        elif token == '<sub>':
            active_sub = True
        elif token == '</sub>':
            active_sub = False
        elif token == '<sup>':
            active_super = True
        elif token == '</sup>':
            active_super = False
        elif token == '<normal>':
            active_normal = True
        elif token == '</normal>':
            active_normal = False
        elif token.startswith('<') and token.endswith('>'):
            pass
        else:
            run = p.add_run(token)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.subscript = active_sub
            run.font.superscript = active_super
            run.font.bold = active_bold
            if active_italic and not active_normal:
                run.font.italic = True

def add_equation_block(doc: Any, math_content: str, eq_num: int, is_double_column: bool) -> None:
    """Appends an equation paragraph aligned via tab stops."""
    p = doc.add_paragraph()
    if is_double_column:
        center_tab = Inches(1.69)
        right_tab = Inches(3.38)
    else:
        center_tab = Inches(3.25)
        right_tab = Inches(6.5)
        
    p.paragraph_format.tab_stops.add_tab_stop(center_tab, WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    
    p.add_run("\t")
    
    tagged_math = latex_math_to_tagged_string(math_content)
    add_math_to_paragraph(p, tagged_math)
    
    p.add_run(f"\t({eq_num})")

def add_figure_block(doc: Any, fig_num: int, img_ref: str, caption: str, is_double_column: bool, is_star_env: bool, label_map: Dict[str, Any], citation_map: Dict[str, int]) -> None:
    """Appends a centered placeholder box with generous vertical space for the figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    
    # Render placeholder layout depending on the target column width
    if is_double_column and not is_star_env:
        box_text = (
            "\n"
            "┌──────────────────────────────┐\n"
            f"│  [IMAGE PLACEHOLDER: Fig. {fig_num}] │\n"
            f"│  Ref: {img_ref[:18]:<18} │\n"
            "│                              │\n"
            "│  (Insert/paste picture here) │\n"
            "└──────────────────────────────┘\n"
            "\n"
        )
    else:
        box_text = (
            "\n"
            "┌────────────────────────────────────────────────────────┐\n"
            "│                                                        │\n"
            f"│         [IMAGE PLACEHOLDER: Fig. {fig_num}]              │\n"
            f"│         Reference: {img_ref:<35} │\n"
            "│                                                        │\n"
            "│         (Insert or paste the image here)               │\n"
            "│                                                        │\n"
            "└────────────────────────────────────────────────────────┘\n"
            "\n"
        )
        
    run = p.add_run(box_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)
        
    # Figure caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(12)
    
    run_label = p_cap.add_run(f"Fig. {fig_num}. ")
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(8)
    run_label.font.bold = True
    
    cleaned_caption = clean_latex_text(caption, label_map, citation_map)
    add_formatted_text(p_cap, cleaned_caption, font_size=8)

def add_table_block(doc: Any, table_num: str, title: str, alignments: List[Any], data: List[List[str]], is_double_column: bool, is_star_env: bool, label_map: Dict[str, Any], citation_map: Dict[str, int]) -> None:
    """Appends formatted table in scientific style with no vertical lines."""
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(8)
    p_cap.paragraph_format.space_after = Pt(4)
    p_cap.paragraph_format.keep_with_next = True
    
    run_label = p_cap.add_run(f"TABLE {table_num}\n")
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(8)
    run_label.font.bold = True
    
    run_title = p_cap.add_run(title.upper())
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(8)
    
    rows_cnt = len(data)
    if rows_cnt == 0:
        return
    cols_cnt = max(len(r) for r in data)
    
    table = doc.add_table(rows=rows_cnt, cols=cols_cnt)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    thin_border = {"sz": 4, "val": "single", "color": "000000"}
    thick_border = {"sz": 12, "val": "single", "color": "000000"}
    
    # Calculate dynamic column widths based on maximum character length
    col_char_lens = [0] * cols_cnt
    for row_data in data:
        for c_idx in range(cols_cnt):
            val = row_data[c_idx] if c_idx < len(row_data) else ""
            col_char_lens[c_idx] = max(col_char_lens[c_idx], len(val))
            
    # Assign minimum length of 5 to avoid tiny columns
    col_char_lens = [max(l, 5) for l in col_char_lens]
    total_chars = sum(col_char_lens)
    
    if is_double_column:
        total_width = Inches(7.25) if is_star_env else Inches(3.38)
    else:
        total_width = Inches(7.0) if is_star_env else Inches(6.5)
        
    col_widths = [total_width * (l / total_chars) for l in col_char_lens]
    
    for r_idx, row in enumerate(table.rows):
        row_data = data[r_idx]
        for c_idx, cell in enumerate(row.cells):
            cell_text = row_data[c_idx].strip() if c_idx < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            
            if c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                
            cleaned_text = clean_latex_text(cleaned_text := cell_text, label_map, citation_map)
            add_formatted_text(p, cleaned_text, font_size=8)
            
            is_header = (r_idx == 0)
            is_last = (r_idx == rows_cnt - 1)
            
            if is_header:
                set_cell_border(cell, top=thick_border, bottom=thin_border)
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                for run in p.runs:
                    run.font.bold = True
            elif is_last:
                set_cell_border(cell, bottom=thick_border)
                set_cell_margins(cell, top=60, bottom=60, left=150, right=150)
            else:
                set_cell_border(cell)
                set_cell_margins(cell, top=60, bottom=60, left=150, right=150)
                
            is_total_row = any(term in str(row_data[0]).upper() for term in ["TOTAL", "SUM"]) if row_data else False
            if is_total_row:
                for run in p.runs:
                    run.font.bold = True
                    
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_widths[c_idx]

def parse_tabular(tabular_text: str) -> Tuple[List[Any], List[List[str]]]:
    """Parses tabular alignment specifications and row contents cleanly."""
    # Find \begin{tabular}, \begin{tabularx}, or \begin{longtable}
    match = re.search(r'\\begin\{(tabular|tabularx|longtable)\}', tabular_text)
    alignments = []
    start_idx = 0
    if match:
        env_name = match.group(1)
        pos = match.end()
        
        # Skip optional brackets [...]
        while pos < len(tabular_text) and tabular_text[pos].isspace():
            pos += 1
        if pos < len(tabular_text) and tabular_text[pos] == '[':
            _, pos = get_optional_brackets(tabular_text, pos)
            
        # For tabularx, the first braced arg is the width (e.g. {\textwidth}). Skip it to find spec.
        braced_count = 0
        max_braced = 2 if env_name == "tabularx" else 1
        
        spec = ""
        while braced_count < max_braced:
            while pos < len(tabular_text) and tabular_text[pos].isspace():
                pos += 1
            if pos < len(tabular_text) and tabular_text[pos] == '{':
                brace_content, end_pos = get_nested_brackets(tabular_text, pos)
                braced_count += 1
                pos = end_pos
                if braced_count == max_braced:
                    spec = brace_content[1:-1].strip()
                    start_idx = end_pos
            else:
                break
                
        if spec:
            i = 0
            while i < len(spec):
                char = spec[i]
                if char == 'l':
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
                    i += 1
                elif char == 'c':
                    alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                    i += 1
                elif char == 'r':
                    alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                    i += 1
                elif char in ['X', 'L', 'C', 'R', 'p', 'm', 'b']:
                    if char == 'C':
                        alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                    elif char == 'R':
                        alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                    else:
                        alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
                    i += 1
                    # Skip braces content (e.g. width specification p{2cm})
                    while i < len(spec) and spec[i].isspace():
                        i += 1
                    if i < len(spec) and spec[i] == '{':
                        _, end_p = get_nested_brackets(spec, i)
                        i = end_p
                else:
                    i += 1
        else:
            start_idx = tabular_text.find('}') + 1
    else:
        start_idx = tabular_text.find('}') + 1
        
    end_match = re.search(r'\\end\{(tabular|tabularx|longtable)\}', tabular_text)
    end_idx = end_match.start() if end_match else len(tabular_text)
    
    content = tabular_text[start_idx:end_idx].strip()
    
    # Character-by-character scan to replace \\ inside braces
    processed_content = []
    brace_depth = 0
    i = 0
    while i < len(content):
        char = content[i]
        if char == '{':
            brace_depth += 1
            processed_content.append(char)
            i += 1
        elif char == '}':
            brace_depth = max(0, brace_depth - 1)
            processed_content.append(char)
            i += 1
        elif content[i:].startswith(r'\\'):
            if brace_depth > 0:
                # Inside braces: replace with placeholder
                processed_content.append('TEMP_CELL_BREAK')
            else:
                # Outside braces: keep as row separator
                processed_content.append(r'\\')
            i += 2
        else:
            processed_content.append(char)
            i += 1
            
    content = "".join(processed_content)
    
    rows = []
    lines = content.split(r'\\')
    for line in lines:
        line = line.strip()
        clean_line = re.sub(r'\\(hline|toprule|midrule|bottomrule|cline\{[^}]+\})', '', line).strip()
        if not clean_line:
            continue
        # Protect escaped ampersands \& from splitting
        clean_line = clean_line.replace(r'\&', 'TEMP_AMPERSAND')
        cells = [c.replace('TEMP_CELL_BREAK', '\n').replace('TEMP_AMPERSAND', r'\&').strip() for c in clean_line.split('&')]
        rows.append(cells)
        
    return alignments, rows

# Sequential block parsing and helpers
def find_matching_end(text: str, env_name: str, start_pos: int) -> int:
    begin_tag = f'\\begin{{{env_name}}}'
    end_tag = f'\\end{{{env_name}}}'
    
    depth = 0
    pos = start_pos
    
    while pos < len(text):
        next_begin = text.find(begin_tag, pos)
        next_end = text.find(end_tag, pos)
        
        if next_end == -1:
            return -1
            
        if next_begin != -1 and next_begin < next_end:
            depth += 1
            pos = next_begin + len(begin_tag)
        else:
            depth -= 1
            if depth == 0:
                return next_end
            pos = next_end + len(end_tag)
            
    return -1

def parse_latex_to_blocks(body_text: str) -> List[Dict[str, Any]]:
    """Scans body text sequentially and separates it into structural nodes."""
    blocks = []
    pos = 0
    
    while pos < len(body_text):
        while pos < len(body_text) and body_text[pos].isspace():
            pos += 1
        if pos >= len(body_text):
            break
            
        if body_text[pos:].startswith(r'\begin{'):
            env_match = re.match(r'\\begin\{([^}]+)\}', body_text[pos:])
            if env_match:
                env_name = env_match.group(1).strip()
                end_tag = f'\\end{{{env_name}}}'
                end_idx = find_matching_end(body_text, env_name, pos)
                if end_idx != -1:
                    env_content = body_text[pos + env_match.end() : end_idx].strip()
                    blocks.append({
                        "type": "environment",
                        "name": env_name,
                        "content": env_content
                    })
                    pos = end_idx + len(end_tag)
                    continue
                else:
                    end_idx = body_text.find(end_tag, pos)
                    if end_idx != -1:
                        env_content = body_text[pos + env_match.end() : end_idx].strip()
                        blocks.append({
                            "type": "environment",
                            "name": env_name,
                            "content": env_content
                        })
                        pos = end_idx + len(end_tag)
                        continue
                        
        sec_match = re.match(r'\\(section|subsection|subsubsection|paragraph)\*?\s*\{', body_text[pos:])
        if sec_match:
            sec_type = sec_match.group(1)
            brace_start = pos + sec_match.end() - 1
            content, end_pos = get_nested_brackets(body_text, brace_start)
            blocks.append({
                "type": "heading",
                "level": sec_type,
                "title": content[1:-1].strip()
            })
            pos = end_pos
            continue
            
        next_pos = len(body_text)
        para_break = body_text.find('\n\n', pos)
        if para_break != -1 and para_break < next_pos:
            next_pos = para_break
            
        env_begin = body_text.find(r'\begin{', pos + 1)
        if env_begin != -1 and env_begin < next_pos:
            next_pos = env_begin
            
        for h_cmd in [r'\section', r'\subsection', r'\subsubsection']:
            h_pos = body_text.find(h_cmd, pos + 1)
            if h_pos != -1 and h_pos < next_pos:
                next_pos = h_pos
                
        text_chunk = body_text[pos:next_pos].strip()
        if text_chunk:
            clean_test = re.sub(r'%.*?\n', '\n', text_chunk).strip()
            if clean_test:
                blocks.append({
                    "type": "paragraph",
                    "content": text_chunk
                })
        
        # Ensure pos always advances, preventing infinite loops on syntax edge cases
        if pos == next_pos:
            pos += 1
        else:
            pos = next_pos
            
        if para_break != -1 and pos == para_break:
            pos += 2
            
    return blocks

# Numbering conversions
def to_roman(n: int) -> str:
    roman_numerals = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", 
                      "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX"]
    if 0 < n < len(roman_numerals):
        return roman_numerals[n]
    return str(n)

def to_alpha(n: int) -> str:
    if 1 <= n <= 26:
        return chr(64 + n)
    return str(n)

def configure_section_margins(section: Any, top: float, bottom: float, left: float, right: float) -> None:
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def write_buffered_paragraph(doc: Any, text: str, label_map: Dict[str, Any], citation_map: Dict[str, int]) -> None:
    global just_after_heading, just_after_table_fig
    text = text.strip()
    if not text:
        return
        
    # Extract leading label if any
    if text.startswith(r'\label{'):
        lbl_match = re.match(r'\\label\{([^}]+)\}', text)
        if lbl_match:
            text = text[lbl_match.end():].strip()
            
    text = clean_latex_text(text, label_map, citation_map)
    if not text:
        return
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    space_before = 8 if just_after_table_fig else 0
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(0)
    
    first_word = text.split()[0] if text.split() else ""
    actual_first_word = re.sub(r'<[^>]+>', '', first_word).strip()
    is_continuation = False
    if actual_first_word and actual_first_word[0].islower() and actual_first_word not in ['i', 'a']:
        is_continuation = True
        
    if just_after_heading or is_continuation:
        p.paragraph_format.first_line_indent = Inches(0)
    else:
        p.paragraph_format.first_line_indent = Inches(0.15)
        
    p.paragraph_format.line_spacing = Pt(10)
    
    just_after_heading = False
    just_after_table_fig = False
    
    add_formatted_text(p, text)

# Recursive block processing to compile nodes
def process_block(
    doc: Any, 
    block: Dict[str, Any], 
    state: Dict[str, int], 
    is_double_column: bool, 
    is_ieee: bool, 
    label_map: Dict[str, Any], 
    citation_map: Dict[str, int]
) -> None:
    global just_after_heading, just_after_table_fig
    try:
        if block["type"] == "heading":
            level = block["level"]
            title = block["title"]
            title = re.sub(r'\\label\{[^}]+\}', '', title).strip()
            title = clean_latex_text(title, label_map, citation_map)
            
            p_head = doc.add_paragraph()
            p_head.paragraph_format.keep_with_next = True
            
            if level == "section":
                state["sec"] += 1
                state["subsec"] = 0
                state["subsubsec"] = 0
                
                if is_ieee:
                    roman = to_roman(state["sec"])
                    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_head.paragraph_format.space_before = Pt(12)
                    p_head.paragraph_format.space_after = Pt(6)
                    run = p_head.add_run(f"{roman}. {title.upper()}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.bold = True
                else:
                    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_head.paragraph_format.space_before = Pt(14)
                    p_head.paragraph_format.space_after = Pt(6)
                    run = p_head.add_run(f"{state['sec']}. {title}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    
            elif level == "subsection":
                state["subsec"] += 1
                state["subsubsec"] = 0
                
                if is_ieee:
                    alpha = to_alpha(state["subsec"])
                    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_head.paragraph_format.space_before = Pt(8)
                    p_head.paragraph_format.space_after = Pt(4)
                    run = p_head.add_run(f"{alpha}. {title}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.italic = True
                else:
                    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_head.paragraph_format.space_before = Pt(10)
                    p_head.paragraph_format.space_after = Pt(4)
                    run = p_head.add_run(f"{state['sec']}.{state['subsec']} {title}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    
            elif level == "subsubsection":
                state["subsubsec"] += 1
                
                if is_ieee:
                    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_head.paragraph_format.space_before = Pt(6)
                    p_head.paragraph_format.space_after = Pt(2)
                    run = p_head.add_run(f"{state['subsubsec']}) {title}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.italic = True
                else:
                    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_head.paragraph_format.space_before = Pt(8)
                    p_head.paragraph_format.space_after = Pt(2)
                    run = p_head.add_run(f"{state['sec']}.{state['subsec']}.{state['subsubsec']} {title}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.bold = True
                    
            just_after_heading = True
            
        elif block["type"] == "paragraph":
            write_buffered_paragraph(doc, block["content"], label_map, citation_map)
            
        elif block["type"] == "environment":
            name = block["name"]
            content = block["content"]
            is_star = name.endswith('*')
            base_name = name.rstrip('*')
            
            if base_name == "itemize" or base_name == "enumerate":
                items = re.split(r'\\item', content)
                for it_idx, item in enumerate(items):
                    item = item.strip()
                    if not item:
                        continue
                    
                    # Recursively parse item content to handle nested environments (like equation blocks)
                    nested_blocks = parse_latex_to_blocks(item)
                    for n_idx, n_block in enumerate(nested_blocks):
                        if n_block["type"] == "paragraph":
                            p_item = doc.add_paragraph()
                            p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_item.paragraph_format.left_indent = Inches(0.2)
                            p_item.paragraph_format.first_line_indent = Inches(-0.1)
                            p_item.paragraph_format.space_after = Pt(2)
                            
                            # Add bullet prefix only to the first paragraph of the list item
                            if n_idx == 0:
                                if base_name == "itemize":
                                    p_item.add_run("• ").font.name = 'Times New Roman'
                                else:
                                    p_item.add_run(f"{it_idx}. ").font.name = 'Times New Roman'
                            
                            cleaned_item = clean_latex_text(n_block["content"], label_map, citation_map)
                            add_formatted_text(p_item, cleaned_item)
                        else:
                            # Process nested environments recursively (e.g. equation blocks)
                            process_block(doc, n_block, state, is_double_column, is_ieee, label_map, citation_map)
                    
            elif base_name == "equation":
                state["eq"] += 1
                math_clean = re.sub(r'\\label\{[^}]+\}', '', content).strip()
                add_equation_block(doc, math_clean, state["eq"], is_double_column)
                
            elif base_name == "align":
                lines = content.split(r'\\')
                for line in lines:
                    if line.strip():
                        state["eq"] += 1
                        math_clean = re.sub(r'\\label\{[^}]+\}', '', line).strip()
                        add_equation_block(doc, math_clean, state["eq"], is_double_column)
                        
            elif base_name == "figure":
                state["fig"] += 1
                img_match = re.search(r'\\includegraphics\s*(?:\[[^\]]*\])?\s*\{([^}]+)\}', content)
                img_ref = img_match.group(1).strip() if img_match else "unknown_image"
                
                cap_match = re.search(r'\\caption\s*\{', content)
                caption = ""
                if cap_match:
                    cap_start = cap_match.end() - 1
                    cap_brace, _ = get_nested_brackets(content, cap_start)
                    caption = cap_brace[1:-1].strip()
                    caption = re.sub(r'\\label\{[^}]+\}', '', caption).strip()
                    
                add_figure_block(doc, state["fig"], img_ref, caption, is_double_column, is_star, label_map, citation_map)
                just_after_table_fig = True
                
            elif base_name == "table":
                state["tbl"] += 1
                cap_match = re.search(r'\\caption\s*\{', content)
                title = "Table Title"
                if cap_match:
                    cap_start = cap_match.end() - 1
                    cap_brace, _ = get_nested_brackets(content, cap_start)
                    title = cap_brace[1:-1].strip()
                    title = re.sub(r'\\label\{[^}]+\}', '', title).strip()
                    
                tabular_match = re.search(r'\\begin\{(tabular|tabularx|longtable)\}.*?\\end\{\1\}', content, re.DOTALL)
                if tabular_match:
                    alignments, data = parse_tabular(tabular_match.group(0))
                    roman_tbl = to_roman(state["tbl"])
                    add_table_block(doc, roman_tbl, title, alignments, data, is_double_column, is_star, label_map, citation_map)
                    just_after_table_fig = True
                else:
                    print(f"Warning: tabular/tabularx/longtable environment not found inside table block {state['tbl']}")
            else:
                # Generic unknown environment (e.g. proof, theorem, verbatim, center, quote)
                # Parse its content recursively and append those blocks!
                print(f"Generic/custom environment parsed: {name}")
                nested_blocks = parse_latex_to_blocks(content)
                for n_block in nested_blocks:
                    process_block(doc, n_block, state, is_double_column, is_ieee, label_map, citation_map)
    except Exception as e:
        print(f"Warning: Failed to process block {block.get('type')}: {e}")
        try:
            if block.get("type") == "paragraph":
                p = doc.add_paragraph()
                p.add_run(str(block.get("content", "")))
        except:
            pass

def main() -> None:
    global just_after_heading, just_after_table_fig
    
    # Command line inputs
    if len(sys.argv) < 2:
        tex_path = "d:/IOT paper/IOT_EL/main.tex"
        out_path = "d:/IOT paper/IOT_EL/crowd_safety_oracle_ieee.docx"
        print(f"No arguments provided. Running default compilation: {tex_path} -> {out_path}")
    else:
        tex_path = sys.argv[1]
        if len(sys.argv) >= 3:
            out_path = sys.argv[2]
        else:
            base, _ = os.path.splitext(tex_path)
            out_path = base + ".docx"
            
    if not os.path.exists(tex_path):
        print(f"Error: LaTeX source file not found at {tex_path}")
        return
        
    with open(tex_path, 'r', encoding='utf-8') as f:
        tex_content = f.read()
        
    # Run pre-processing to expand custom resume environments/macros and basic symbols
    tex_content = preprocess_latex(tex_content)
        
    # Detect document layout style
    is_double_column = False
    is_ieee = False
    class_match = re.search(r'\\documentclass\s*(?:\[([^\]]*)\])?\s*\{([^}]+)\}', tex_content)
    if class_match:
        options = class_match.group(1) or ""
        doc_class = class_match.group(2).strip()
        print(f"Document class detected: {doc_class} (Options: {options})")
        if doc_class == "IEEEtran":
            is_double_column = True
            is_ieee = True
        elif "twocolumn" in options:
            is_double_column = True
            
    print(f"Layout mode: {'Double-column' if is_double_column else 'Single-column'}")
    
    # Extract Title and optional Thanks
    title_raw, _ = extract_command_arg(tex_content, r"\title")
    thanks_text = ""
    if title_raw:
        thanks_match = re.search(r'\\thanks\s*\{', title_raw)
        if thanks_match:
            th_start = thanks_match.end() - 1
            th_brace, _ = get_nested_brackets(title_raw, th_start)
            thanks_text = th_brace[1:-1].strip()
            title_text = title_raw[:thanks_match.start()] + title_raw[thanks_match.end() + len(th_brace) - 2:]
            title_text = title_text.strip()
        else:
            title_text = title_raw
    else:
        # Fallback for resumes: Try to find name in {\Huge \scshape Name} or similar header blocks
        name_match = re.search(r'\{\\(Huge|huge|Large|large)\s*(?:\\scshape\s*|\\bfseries\s*)*([^{}]+)\}', tex_content, re.IGNORECASE)
        if name_match:
            title_text = name_match.group(2).strip()
            print(f"Extracted document title from header: {title_text}")
        else:
            title_text = "Untitled Document"
        
    # Extract Authors
    author_raw, _ = extract_command_arg(tex_content, r"\author")
    authors = []
    if author_raw:
        n_blocks = re.findall(r'\\IEEEauthorblockN\s*\{', author_raw)
        if n_blocks:
            pos = 0
            while True:
                match_n = re.search(r'\\IEEEauthorblockN\s*\{', author_raw[pos:])
                if not match_n:
                    break
                n_start = pos + match_n.end() - 1
                name_brace, n_end = get_nested_brackets(author_raw, n_start)
                name = name_brace[1:-1].strip()
                
                match_a = re.search(r'\\IEEEauthorblockA\s*\{', author_raw[n_end:])
                if match_a:
                    a_start = n_end + match_a.end() - 1
                    aff_brace, a_end = get_nested_brackets(author_raw, a_start)
                    aff = aff_brace[1:-1].strip()
                    pos = a_end
                else:
                    aff = ""
                    pos = n_end
                authors.append({"name": name, "affiliation": aff})
        else:
            parts = author_raw.split(r'\and')
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                lines = [l.strip() for l in part.split(r'\\')]
                name = lines[0]
                aff = "\n".join(lines[1:])
                authors.append({"name": name, "affiliation": aff})
                
    # Extract bibliography
    citations = []
    citation_map = {}
    bib_match = re.search(r'\\begin\{thebibliography\}(.*?)\\end\{thebibliography\}', tex_content, re.DOTALL)
    if bib_match:
        bib_content = bib_match.group(1)
        bib_items_raw = re.split(r'\\bibitem\s*(?:\[[^\]]*\])?\s*\{([^}]+)\}', bib_content)
        counter = 1
        for idx in range(1, len(bib_items_raw), 2):
            label = bib_items_raw[idx].strip()
            text = bib_items_raw[idx+1].strip()
            citations.append((label, text, counter))
            citation_map[label] = counter
            counter += 1
            
    # Extract body content (between \begin{document} and \end{document})
    doc_match = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', tex_content, re.DOTALL)
    if doc_match:
        body_text = doc_match.group(1)
    else:
        body_text = tex_content
        
    # Strip document header commands from body text if they were placed inside the document environment
    body_text = re.sub(r'\\maketitle', '', body_text)
    body_text = strip_latex_command(body_text, r"\title")
    body_text = strip_latex_command(body_text, r"\author")
    body_text = strip_latex_command(body_text, r"\date")
    body_text = strip_latex_command(body_text, r"\thanks")
        
    # Extract abstract and keywords
    abstract_text = ""
    abs_match = re.search(r'\\begin\{abstract\}(.*?)\\end\{abstract\}', body_text, re.DOTALL)
    if abs_match:
        abstract_text = abs_match.group(1).strip()
        body_text = body_text.replace(abs_match.group(0), "")
        
    keywords_text = ""
    key_match = re.search(r'\\begin\{IEEEkeywords\}(.*?)\\end\{IEEEkeywords\}', body_text, re.DOTALL)
    if key_match:
        keywords_text = key_match.group(1).strip()
        body_text = body_text.replace(key_match.group(0), "")
        
    # Remove bibliography block from body
    if bib_match:
        body_text = body_text.replace(bib_match.group(0), "")
        
    # Parse body content into sequential blocks
    blocks = parse_latex_to_blocks(body_text)
    print(f"Total document blocks parsed: {len(blocks)}")
    
    # ------------------ PASS 1: Reference Mapping ------------------
    label_map = {}
    sec_count = 0
    subsec_count = 0
    subsubsec_count = 0
    eq_count = 0
    fig_count = 0
    tbl_count = 0
    
    # Standard mapping helper
    def map_labels_in_blocks(block_list: List[Dict[str, Any]]) -> None:
        nonlocal sec_count, subsec_count, subsubsec_count, eq_count, fig_count, tbl_count
        for blk in block_list:
            b_text = str(blk.get("content", "")) or str(blk.get("title", ""))
            
            if blk["type"] == "heading":
                level = blk["level"]
                if level == "section":
                    sec_count += 1
                    subsec_count = 0
                    subsubsec_count = 0
                    num_str = to_roman(sec_count) if is_ieee else str(sec_count)
                elif level == "subsection":
                    subsec_count += 1
                    subsubsec_count = 0
                    num_str = to_alpha(subsec_count) if is_ieee else f"{sec_count}.{subsec_count}"
                elif level == "subsubsection":
                    subsubsec_count += 1
                    num_str = f"{subsubsec_count})" if is_ieee else f"{sec_count}.{subsec_count}.{subsubsec_count}"
                    
                lbl_matches = re.findall(r'\\label\{([^}]+)\}', blk["title"])
                for lbl in lbl_matches:
                    label_map[lbl.strip()] = num_str
                    
            elif blk["type"] == "paragraph":
                lbl_matches = re.findall(r'\\label\{([^}]+)\}', blk["content"])
                for lbl in lbl_matches:
                    curr_sec_num = to_roman(sec_count) if is_ieee else str(sec_count)
                    label_map[lbl.strip()] = curr_sec_num
                    
            elif blk["type"] == "environment":
                name = blk["name"]
                content = blk["content"]
                base_name = name.rstrip('*')
                
                if base_name in ["equation", "align"]:
                    if base_name == "equation":
                        eq_count += 1
                        lbl_matches = re.findall(r'\\label\{([^}]+)\}', content)
                        for lbl in lbl_matches:
                            label_map[lbl.strip()] = eq_count
                    else:
                        lines = content.split(r'\\')
                        for line in lines:
                            if line.strip():
                                eq_count += 1
                                lbl_matches = re.findall(r'\\label\{([^}]+)\}', line)
                                for lbl in lbl_matches:
                                    label_map[lbl.strip()] = eq_count
                                    
                elif base_name == "figure":
                    fig_count += 1
                    lbl_matches = re.findall(r'\\label\{([^}]+)\}', content)
                    for lbl in lbl_matches:
                        label_map[lbl.strip()] = fig_count
                        
                elif base_name == "table":
                    tbl_count += 1
                    roman_tbl = to_roman(tbl_count)
                    lbl_matches = re.findall(r'\\label\{([^}]+)\}', content)
                    for lbl in lbl_matches:
                        label_map[lbl.strip()] = roman_tbl
                else:
                    # Recursive mapping for generic containers
                    nested = parse_latex_to_blocks(content)
                    map_labels_in_blocks(nested)
                    
    map_labels_in_blocks(blocks)
    print(f"Cross-reference label map built with {len(label_map)} entries.")
    
    # ------------------ PASS 2: Document Generation ------------------
    doc = Document()
    
    # Margins for Section 1 (Title/Authors)
    section1 = doc.sections[0]
    if is_double_column:
        configure_section_margins(section1, 0.75, 1.0, 0.625, 0.625)
    else:
        configure_section_margins(section1, 1.0, 1.0, 1.0, 1.0)
        
    # Write Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(12)
    
    cleaned_title = clean_latex_text(title_text, label_map, citation_map)
    add_formatted_text(p_title, cleaned_title, font_size=24)
    for run in p_title.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
    
    # Write Thanks (Funding block)
    if thanks_text:
        p_thanks = doc.add_paragraph()
        p_thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_thanks.paragraph_format.space_after = Pt(18)
        cleaned_thanks = clean_latex_text(thanks_text, label_map, citation_map)
        add_formatted_text(p_thanks, cleaned_thanks, font_size=9)
        for run in p_thanks.runs:
            run.font.name = 'Times New Roman'
            run.font.italic = True
        
    # Write Authors side-by-side
    authors_cnt = len(authors)
    if authors_cnt > 0:
        cols_per_row = min(authors_cnt, 3)
        rows_cnt = (authors_cnt + cols_per_row - 1) // cols_per_row
        
        authors_table = doc.add_table(rows=rows_cnt, cols=cols_per_row)
        authors_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        body_width = Inches(7.25) if is_double_column else Inches(6.5)
        col_width = body_width / cols_per_row
        
        for a_idx, auth in enumerate(authors):
            r_idx = a_idx // cols_per_row
            c_idx = a_idx % cols_per_row
            
            cell = authors_table.rows[r_idx].cells[c_idx]
            cell.width = col_width
            set_cell_border(cell)
            set_cell_margins(cell, top=0, bottom=0, left=50, right=50)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            
            name_clean = clean_latex_text(auth["name"], {}, {})
            aff_clean = clean_latex_text(auth["affiliation"], {}, {})
            
            # Add name using add_formatted_text
            p_start_run_idx = len(p.runs)
            add_formatted_text(p, name_clean, font_size=11)
            for run in p.runs[p_start_run_idx:]:
                run.font.name = 'Times New Roman'
                run.font.bold = True
            p.add_run("\n")
            
            lines = aff_clean.split('\\\\')
            if len(lines) == 1:
                lines = aff_clean.split('\n')
                
            for l_idx, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                is_email = "@" in line or "email" in line.lower()
                
                # Add affiliation line using add_formatted_text
                p_start_run_idx = len(p.runs)
                add_formatted_text(p, line, font_size=10)
                for run in p.runs[p_start_run_idx:]:
                    run.font.name = 'Times New Roman'
                    if not is_email:
                        run.font.italic = True
                if l_idx < len(lines)-1:
                    p.add_run("\n")
                    
    # Double column section division
    if is_double_column:
        section2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
        configure_section_margins(section2, 0.75, 1.0, 0.625, 0.625)
        
        sectPr = section2._sectPr
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '360')
        cols.set(qn('w:equalWidth'), '1')
        sectPr.append(cols)
        
    # Write Abstract
    if abstract_text:
        cleaned_abs = clean_latex_text(abstract_text, label_map, citation_map)
        p_abs = doc.add_paragraph()
        p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_abs.paragraph_format.space_before = Pt(12)
        p_abs.paragraph_format.space_after = Pt(6)
        p_abs.paragraph_format.first_line_indent = Inches(0)
        
        run_h = p_abs.add_run("Abstract---")
        run_h.font.name = 'Times New Roman'
        run_h.font.size = Pt(9)
        run_h.font.bold = True
        run_h.font.italic = True
        
        add_formatted_text(p_abs, cleaned_abs, font_size=9, is_abstract=True)
        
    # Write Keywords
    if keywords_text:
        cleaned_key = clean_latex_text(keywords_text, label_map, citation_map)
        p_key = doc.add_paragraph()
        p_key.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_key.paragraph_format.space_before = Pt(0)
        p_key.paragraph_format.space_after = Pt(12)
        p_key.paragraph_format.first_line_indent = Inches(0)
        
        run_h = p_key.add_run("Index Terms---")
        run_h.font.name = 'Times New Roman'
        run_h.font.size = Pt(9)
        run_h.font.bold = True
        run_h.font.italic = True
        
        add_formatted_text(p_key, cleaned_key, font_size=9, is_abstract=True)
        
    # State tracking dictionary for sequential pass 2
    state = {
        "sec": 0,
        "subsec": 0,
        "subsubsec": 0,
        "eq": 0,
        "fig": 0,
        "tbl": 0
    }
    
    # Process all parsed blocks recursively
    for block in blocks:
        process_block(doc, block, state, is_double_column, is_ieee, label_map, citation_map)
                    
    # Write Bibliography Section
    if citations:
        add_references(doc, citations, label_map, citation_map)
        
    doc.save(out_path)
    print(f"Success: LaTeX successfully converted to Word document: {out_path}")

def add_references(doc: Any, citations: List[Tuple[str, str, int]], label_map: Dict[str, Any], citation_map: Dict[str, int]) -> None:
    """Appends reference list at the end of the document in IEEE / standard format."""
    p_ref_head = doc.add_paragraph()
    p_ref_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ref_head.paragraph_format.space_before = Pt(18)
    p_ref_head.paragraph_format.space_after = Pt(6)
    p_ref_head.paragraph_format.keep_with_next = True
    
    run_ref = p_ref_head.add_run("REFERENCES")
    run_ref.font.name = 'Times New Roman'
    run_ref.font.size = Pt(8)
    run_ref.font.bold = True
    
    for label, text, num in citations:
        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_ref.paragraph_format.left_indent = Inches(0.25)
        p_ref.paragraph_format.first_line_indent = Inches(-0.25)
        p_ref.paragraph_format.space_before = Pt(0)
        p_ref.paragraph_format.space_after = Pt(4)
        
        text = text.replace('``', '"').replace("''", '"')
        cleaned_text = clean_latex_text(text, label_map, citation_map)
        full_ref_text = f"[{num}] {cleaned_text}"
        
        add_formatted_text(p_ref, full_ref_text, font_size=8)

if __name__ == '__main__':
    main()
